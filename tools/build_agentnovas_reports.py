from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "artifacts"
OUTPUT.mkdir(parents=True, exist_ok=True)

NAVY = "0B1F33"
BLUE = "1565C0"
CYAN = "16B8D4"
PALE = "EAF4FB"
LIGHT = "F5F8FB"
GOLD = "C48A2A"
RED = "B33A3A"
GREEN = "187B5B"
INK = "152233"
MUTED = "52677A"
WHITE = "FFFFFF"
LINE = "CFDCE7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    tail = paragraph.add_run(" 页")
    tail.font.size = Pt(8)


def set_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, name: str = "Arial Unicode MS") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_document(title: str, subject: str) -> Document:
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "AgentNovas / Codex"
    doc.core_properties.keywords = "AgentNovas, WOX, 运营后台, 运维后台, 权限, 分成"
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.72)
    section.footer_distance = Cm(0.72)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in (("Title", 28, NAVY), ("Heading 1", 20, NAVY), ("Heading 2", 14, BLUE), ("Heading 3", 11.5, INK)):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("AGENTNOVAS  ·  WOX LOCAL REVIEW")
    set_font(run, 8, True, CYAN)
    hp.paragraph_format.space_after = Pt(2)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.add_run("内部工作文档  ·  2026-08-19  ·  ")
    page_number(fp)
    return doc


def add_cover(doc: Document, title: str, subtitle: str, badge: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("A")
    set_font(r, 34, True, WHITE)
    r._element.get_or_add_rPr().append(_run_shading(CYAN))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("AGENTNOVAS")
    set_font(r, 12, True, CYAN)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run(subtitle)
    set_font(r, 13, False, MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(17.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 180, 220, 180, 220)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(badge)
    set_font(r, 11, True, WHITE)
    p = cell.add_paragraph("基于当前本地目录 /Users/zhonghetong/Documents/Codex/2026-08-11/wox 的代码、数据库结构、接口与页面实测结果。")
    p.paragraph_format.space_before = Pt(5)
    for run in p.runs:
        set_font(run, 9, False, "C9D8E5")

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Cm(4.1)
    meta.columns[1].width = Cm(13.4)
    rows = [
        ("审阅日期", "2026 年 8 月 19 日"),
        ("项目状态", "本地 WOX 已构建通过；未推送 Git，未上传服务器"),
        ("当前域名规划", "上传对象暂为 www.tzxsea.com；www.agentnovas.com 尚待 DNS 绑定"),
        ("保密级别", "内部使用 / 运营、技术与管理层参考"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_no_split(row)
        set_cell_shading(row.cells[0], PALE)
        for cell in row.cells:
            set_cell_margins(cell)
        rr = row.cells[0].paragraphs[0].add_run(label)
        set_font(rr, 9, True, BLUE)
        rv = row.cells[1].paragraphs[0].add_run(value)
        set_font(rv, 9, False, INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("AgentNovas · 智能交易中枢")
    set_font(r, 9, True, MUTED)
    doc.add_page_break()


def _run_shading(fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def add_toc(doc: Document, entries: Sequence[tuple[str, str]]) -> None:
    doc.add_heading("文档导航", level=1)
    p = doc.add_paragraph("本目录为静态导航，便于在 WPS 中快速阅读。")
    for number, title in entries:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(1.6)
        table.columns[1].width = Cm(15.9)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_cell_shading(table.cell(0, 0), NAVY)
        set_cell_shading(table.cell(0, 1), LIGHT)
        set_cell_margins(table.cell(0, 0), 80, 100, 80, 100)
        set_cell_margins(table.cell(0, 1), 80, 120, 80, 120)
        r = table.cell(0, 0).paragraphs[0].add_run(number)
        set_font(r, 9, True, WHITE)
        r = table.cell(0, 1).paragraphs[0].add_run(title)
        set_font(r, 9.5, True, INK)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_page_break()


def add_callout(doc: Document, title: str, body: str, color: str = CYAN) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(0.35)
    table.columns[1].width = Cm(17.15)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_shading(table.cell(0, 0), color)
    set_cell_shading(table.cell(0, 1), LIGHT)
    set_cell_margins(table.cell(0, 0), 80, 40, 80, 40)
    set_cell_margins(table.cell(0, 1), 120, 150, 120, 150)
    p = table.cell(0, 1).paragraphs[0]
    r = p.add_run(title)
    set_font(r, 10, True, NAVY)
    p = table.cell(0, 1).add_paragraph(body)
    p.paragraph_format.space_before = Pt(3)
    for run in p.runs:
        set_font(run, 9, False, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: Iterable[str], level: int = 0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.left_indent = Cm(0.5 + level * 0.45)
        p.paragraph_format.first_line_indent = Cm(-0.2)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_font(r, 9.5, False, INK)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, 9.5, False, INK)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None, font_size: float = 8.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Cm(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_font(r, font_size, True, WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        set_no_split(row)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, font_size, idx == 0, INK if idx == 0 else MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc: Document, number: str, title: str, intro: str | None = None) -> None:
    p = doc.add_heading(level=1)
    r = p.add_run(f"{number}  {title}")
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(8)


def subsection(doc: Document, title: str, body: str | None = None) -> None:
    doc.add_heading(title, level=2)
    if body:
        doc.add_paragraph(body)


def add_status_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    add_table(doc, ["检查项", "结论", "说明"], rows, [4.1, 2.3, 11.1], 8.7)


def build_website_guide() -> Path:
    doc = setup_document("AgentNovas 网站功能与后台体系说明书", "网站功能、后台体系、权限范围与奖励分成说明")
    add_cover(doc, "网站功能与后台体系说明书", "核心功能 · 运营后台 · 运维后台 · 权限矩阵 · 奖励分成", "当前结论：产品骨架与后台治理体系已形成，构建与核心页面验收通过")
    add_toc(doc, [
        ("01", "执行摘要与产品定位"), ("02", "全站架构与数据边界"), ("03", "用户侧核心功能"),
        ("04", "运营后台功能体系"), ("05", "运维后台功能体系"), ("06", "角色与数据可见范围"),
        ("07", "客户 360° 详情"), ("08", "行情、新闻与交易成熟度"), ("09", "奖励分成规则"),
        ("10", "核心亮点与竞争优势"), ("11", "本地验收结果"), ("A", "系统规模附录"),
    ])

    heading(doc, "01", "执行摘要与产品定位", "AgentNovas 是一个面向交易用户、组织运营团队和总部技术团队的一体化智能交易平台。当前本地版已完成用户门户、行情与新闻、AI 助手、策略广场、交易账户连接、会员、通知、组织运营和技术运维的统一产品框架。")
    add_callout(doc, "一句话定位", "以客户归属链和权限隔离为基础，把行情研究、AI 策略、交易连接、会员计费、组织运营、策略审核与系统运维放在同一个可审计平台中。")
    subsection(doc, "当前完成度判断")
    add_status_table(doc, [
        ("产品与界面框架", "已形成", "前台、运营后台、运维后台入口和主要页面已具备"),
        ("组织与客户权限", "可用", "五级组织链、直客邀请、客户范围过滤、手动启停和恢复已实现"),
        ("行情与新闻", "部分生产可用", "币安公开行情与 RSS 新闻本地实测可返回真实数据；供应商稳定性与授权仍需加强"),
        ("策略与审核", "核心流程可用", "用户提交、双人策略审核、上架状态与审计链路已具备"),
        ("真实资金与自动交易", "尚未生产开放", "当前仍以模拟执行、接口连接和受控自动化为主"),
        ("运维治理", "本次已增强", "新增持久化系统、功能、计费、集成、安全配置和运行状态面板"),
    ])

    heading(doc, "02", "全站架构与数据边界")
    add_table(doc, ["层级", "主要内容", "责任边界"], [
        ("用户门户", "交易大厅、Agent 对话、行情中心、策略广场、交易中心、会员中心、消息订阅", "客户仅查看与操作自身账户和自身策略"),
        ("运营后台", "待审批、客户管理、组织成员、数据中心、月度分红、结算、收入调整、邀请码", "按照员工→主管→经理→分公司→总公司的组织链控制可见范围"),
        ("运维后台", "策略审核、系统配置、功能开关、AI、计费、权限、集成、系统运维、安全", "暂由 hq_admin 承载，总公司技术人员使用"),
        ("服务端与数据库", "82 个 API 路由、39 个数据表、会话、审计、收入账本、审批单", "接口再次校验角色和归属，不能只依靠前端隐藏"),
        ("外部服务", "行情供应商、RSS 新闻、OpenAI 兼容大模型、交易所连接、通知渠道", "密钥加密、超时保护、运行时路由；生产 SLA 仍需建设"),
    ], [2.6, 8.0, 6.9])
    subsection(doc, "关键数据边界")
    add_bullets(doc, [
        "客户归属由 branchId、managerId、supervisorId、employeeId 构成；服务端按角色逐级过滤。",
        "总公司与总公司客服可查看全站客户；分公司与财务/审核员仅看同分公司；经理、主管、员工只看自己的直属范围。",
        "客户风控信息对运营侧只展示客户本人是否点击过“停止交易”及时间记录，不开放代替客户操作的按钮。",
        "月度分红接口已收紧为当前账号/当前组织的已确认分配，不返回其他组织和总公司的收入汇总。",
        "系统配置写入 platform_settings 表并进入审计日志，前端默认值只作为无配置时的安全回退。",
    ])

    heading(doc, "03", "用户侧核心功能")
    add_table(doc, ["模块", "当前功能", "数据性质 / 备注"], [
        ("交易大厅", "平台概览、Agent 状态、策略与行情入口", "部分展示指标仍为演示值，需逐步替换为实时统计"),
        ("Agent 对话", "市场研究、策略讨论、团队协作、用户自带模型", "支持系统模型和自定义 OpenAI 兼容端点"),
        ("行情中心", "自选、K 线、缩放/拖动、盘口式指标、新闻、多个公开源", "币安实测可用；实时推送与 REST 回补并行"),
        ("策略广场", "平台策略、用户策略、详情、回测、关注与跟随", "用户策略必须先提交并经过总部双人审核"),
        ("交易中心", "交易所 API 账户、路由、持仓、模拟订单与风控", "真实自动下单默认关闭，不能视为生产交易系统"),
        ("会员中心", "套餐、续费、权益、积分与收益费率规则", "支付链路目前仍需真实地址、链上确认和对账服务"),
        ("消息订阅", "站内、邮件、Telegram、WhatsApp 偏好与强制通知", "队列与偏好已具备，外部发送器仍需生产接入"),
        ("登录注册", "邀请码注册、邀请链接自动填码、手机号必填、账号找回", "邀请注册与密码强度可由运维配置控制"),
    ], [2.7, 8.0, 6.8])

    heading(doc, "04", "运营后台功能体系", "运营后台给公司内部人员使用，与运维后台共用数据库，但不拥有系统接口、密钥、安全策略和收款地址的配置权。")
    add_table(doc, ["板块", "谁能看到", "主要内容与操作"], [
        ("待审批", "仅分公司账号", "处理本分公司由下级提交的客户归属、转移、结算、调整和组织关系申请；申请人不能审批自己"),
        ("运营概览", "仅分公司账号", "展示分公司范围内业务总览，不放在员工、主管、经理后台"),
        ("总数据中心", "内部运营角色按范围", "客户、交易、持仓、盈亏、跟随、账户连接和趋势汇总"),
        ("组织成员", "员工不显示", "总公司建分公司、分公司建经理、经理建主管、主管建员工；关系树逐层查看到用户"),
        ("客户管理", "所有内部角色按范围", "搜索筛选、客户 360° 详情、备注、冻结/恢复；高级动作受角色限制"),
        ("结算付款", "分公司/财务/审核角色接口", "结算申请进入审批流程；收款地址不属于运营配置权"),
        ("收入调整", "分公司/财务/审核角色接口", "必须关联客户、来源、金额、原因和证据，审批后进入收入账本"),
        ("月度分红", "分公司、已启用个人代理", "仅展示本组织或本人可见的已确认分配与个人代理阶梯"),
        ("邀请码", "总公司、分公司、经理、主管、员工", "各级邀请码只邀请直客；创建下一级人员仅在组织成员板块完成"),
    ], [2.6, 4.3, 10.6])
    add_callout(doc, "已按要求隐藏", "团队客户任务、月度任务指标、逾期应收款、收款地址和策略跟随权限均不作为当前运营后台常驻板块；其中旧接口/旧组件仍可能保留在代码中，后续可做安全下线。", GOLD)

    heading(doc, "05", "运维后台功能体系", "运维后台以总公司技术部门为目标用户。本次已经把原来的静态说明升级为可读取、可保存、可审计的控制面板，并把策略广场用户上传审核放在第一位。")
    add_table(doc, ["顺序", "板块", "现有能力", "下一层能力"], [
        ("1", "策略审核", "用户上传策略待审提醒、双人审核进度、通过/驳回、版本一致性检查", "审核 SLA、风险评分、代码沙箱、自动合规预检"),
        ("2", "系统配置", "品牌名、域名、客服、版权、默认语言、启用语言、维护公告、预览", "多站点、域名证书状态、翻译版本发布"),
        ("3", "功能开关", "行情、新闻、AI、策略、交易、会员、通知、邀请、自动交易、发布通道", "全部开关统一服务端拦截、灰度人群、回滚版本"),
        ("4", "AI 助手运营", "系统模型、供应商密钥、自定义端点、连接测试、技能状态", "提示词版本库、token 账单、预算与异常告警"),
        ("5", "计费与支付", "结算币种、积分、优惠券、退款、结算日、分成规则展示", "真实支付通道、退款工单、链上确认、自动对账"),
        ("6", "权限与角色", "五级运营矩阵、职能角色边界、服务端授权说明", "独立技术角色、可编辑权限包、临时授权与审批"),
        ("7", "集成管理", "默认行情源、RSS、刷新周期、超时、自定义大模型接口、17 类连接器", "健康探测、熔断、配额、账单、密钥轮换"),
        ("8", "系统运维", "用户/组织/会话/审批/策略/交易/审计统计、健康接口、最近审计", "任务队列、版本发布、备份恢复、回滚、告警中心"),
        ("9", "安全设置", "会话数、密码长度、邮箱验证、IP 审计、限流总开关、紧急停单、IP 名单", "CIDR、代理 IP 规范化、WAF、MFA、密钥托管"),
    ], [1.0, 2.6, 7.6, 6.3], 7.8)
    subsection(doc, "本次运维增强的服务端落点")
    add_bullets(doc, [
        "新增 platform_settings 持久化表、迁移脚本、读写 API 和管理员审计日志。",
        "登录、注册、密码、会话、邮箱验证、IP 名单和平台紧急停止开始读取运维配置。",
        "行情与新闻接口读取默认供应商、RSS、请求超时和刷新周期；默认公共源调整为 Binance。",
        "自动策略周期读取自动交易总开关与紧急停止状态，未开启时不能生成新的自动执行。",
        "运维概览从数据库实时统计，不使用纯静态卡片冒充运行状态。",
    ])

    heading(doc, "06", "角色与数据可见范围")
    add_table(doc, ["能力", "员工", "主管", "经理", "分公司", "总公司"], [
        ("客户范围", "本人直客", "本人及直属员工客户", "本经理链路客户", "全分公司客户", "全站客户"),
        ("客户查询", "查看/搜索", "查看/搜索", "查看/搜索", "查看/搜索/完整管理", "全站查看"),
        ("冻结/恢复客户", "无", "权限范围内", "权限范围内", "权限范围内并可归档", "审计/全局治理"),
        ("组织成员", "不显示", "创建员工", "创建主管", "创建经理", "创建分公司"),
        ("直客邀请码", "可创建", "可创建", "可创建", "可创建", "可创建/公共池码"),
        ("待审批", "提交", "提交", "提交", "最终审核", "策略/全局审计"),
        ("运营概览", "不显示", "不显示", "不显示", "显示", "运维统计"),
        ("月度分红", "仅个人代理", "不显示", "不显示", "显示本分公司", "不在运营端显示"),
    ], [3.1, 2.75, 2.75, 2.75, 3.0, 3.15], 7.5)
    subsection(doc, "职能角色")
    add_table(doc, ["角色", "数据范围", "限制"], [
        ("总公司客服 hq_support", "全站客户与公共客户池", "不能进入技术运维配置，不能查看密钥"),
        ("财务 finance", "所属分公司账本与结算", "不管理组织，不修改系统接口"),
        ("审核员 auditor", "所属分公司审计；总部审核员可参与策略审核", "不直接修改资金或执行配置"),
        ("客户 customer", "仅自身账户、策略、交易与风险状态", "不能查看任何组织或他人数据"),
    ], [4.0, 6.4, 7.1])

    heading(doc, "07", "客户 360° 详情")
    add_table(doc, ["信息组", "已安排字段"], [
        ("基础与价值", "注册时间、状态、VIP 等级、积分余额、累计充值、累计消费"),
        ("账号资料", "注册 IP、用户名、昵称、手机号、邮箱、Telegram、WhatsApp、语言、时区"),
        ("登录安全", "登录类型、首次/最后登录 IP、最后登录时间、设备、浏览器 / User-Agent"),
        ("客户归属", "来源、归属状态、生效时间、直客负责人、分公司、经理、主管、员工"),
        ("组织链路", "总公司 → 分公司 → 经理 → 主管 → 员工 → 客户"),
        ("券商与资金", "券商类型、环境、连接状态、读写权限、提现授权、账户余额/资产结构"),
        ("交易与风控", "订单数、持仓、已实现盈亏、胜率、回撤、手续费、策略跟随"),
        ("客户自助停单记录", "是否由客户本人点击过停止交易、点击时间与历史记录；运营人员不能代点"),
    ], [4.0, 13.5])
    add_callout(doc, "隐私原则", "客户详情按组织链逐级过滤；邮箱列表可脱敏，详情页在服务端二次校验权限。登录与设备信息属于敏感数据，生产环境应增加查看审计和最小权限。", GREEN)

    heading(doc, "08", "行情、新闻与交易成熟度")
    add_status_table(doc, [
        ("币安公开行情", "本地实测成功", "BTC 报价与 K 线返回实时公开数据；默认源已调整为 Binance"),
        ("Coinbase 公开行情", "本地超时", "当前网络下多次 503，应只作为可选源并增加自动熔断"),
        ("K 线交互", "已具备", "支持周期切换、缩放、左右拖动、历史回看、实时成交更新与 REST 回补"),
        ("新闻 RSS", "本地实测成功", "Cointelegraph 等 RSS 可返回最新新闻；刷新周期由运维配置"),
        ("亚洲股票", "尚未完成", "韩国、日本、香港、中国实时股票和 K 线需要合规数据供应商与交易所授权"),
        ("真实交易执行", "默认关闭", "当前平台 AI 周期与交易中心不等于生产实盘路由；上线前必须完成沙箱与风控验收"),
    ])
    subsection(doc, "关于“0.1 秒刷新”")
    add_bullets(doc, [
        "WebSocket 成交推送可接近实时，但不应固定每 0.1 秒重复请求 REST 接口，否则会触发限频、增加成本并造成假实时。",
        "正确做法是：WebSocket 接收逐笔成交并更新当前蜡烛，REST 按 15–60 秒或断线时回补，服务端缓存并统一时间戳。",
        "价格标准需同时明确交易所、交易对、计价币、时区、K 线边界和聚合规则；不同供应商价格不可能完全一致。",
    ])

    heading(doc, "09", "奖励分成规则", "以下为当前服务端业务规则。所有金额均应以“确认到账/确认收入事件”为前提，未确认、失败、撤销不产生可分配收益。")
    subsection(doc, "9.1 会员充值 / 会员收入")
    add_callout(doc, "公式", "有有效归属时：总公司最终 60%，分公司最终 40%。计算过程为：先将毛收入的 50%留在总公司作为运营成本；剩余 50%网站可分配收益再按总公司 20%、分公司 80%分配。无有效客户归属时，100%归总公司。", BLUE)
    add_table(doc, ["毛收入示例", "运营成本（总部）", "网站收益总部 20%", "网站收益分公司 80%", "最终结果"], [
        ("100 USDT", "50 USDT", "10 USDT", "40 USDT", "总部 60 / 分公司 40"),
        ("1,000 USDT", "500 USDT", "100 USDT", "400 USDT", "总部 600 / 分公司 400"),
    ], [2.8, 3.5, 3.5, 3.7, 4.0])
    subsection(doc, "9.2 总公司部门内部比例")
    add_table(doc, ["部门", "占网站可分配收益", "说明"], [
        ("技术部", "2.5%", "属于总公司网站收益内部口径"),
        ("招商部", "2.5%", "属于总公司网站收益内部口径"),
        ("运营部", "15%", "属于总公司网站收益内部口径"),
        ("合计", "20%", "等于网站可分配收益中的总公司份额"),
    ], [4.2, 4.4, 8.9])
    subsection(doc, "9.3 每周盈利分成费率")
    add_table(doc, ["会员周期", "费率", "计费基数"], [
        ("月度 / 普通", "20%", "当周正的已实现净利润"),
        ("季度", "19%", "当周正的已实现净利润"),
        ("年度", "18%", "当周正的已实现净利润"),
        ("终身", "16%", "当周正的已实现净利润"),
    ], [5.0, 3.2, 9.3])
    doc.add_paragraph("亏损周不收取盈利分成，当前规则不设置亏损结转或高水位追补。")
    subsection(doc, "9.4 个人代理月度阶梯（每月清零）")
    add_table(doc, ["当月业绩 USDT", "分成比例"], [
        ("< 1,000", "20%"), ("1,000–4,999.99", "25%"), ("5,000–9,999.99", "30%"),
        ("10,000–19,999.99", "35%"), ("20,000–49,999.99", "40%"), ("≥ 50,000", "50%"),
    ], [9.5, 8.0])
    subsection(doc, "9.5 策略广场分成")
    add_bullets(doc, [
        "已确认收款的策略费用：平台 50%，策略作者 50%。",
        "有有效客户归属时，平台获得的 50%再按总部 20% / 分公司 80%分配；折算总策略费为总部 10%、分公司 40%、作者 50%。",
        "无有效归属时，平台的 50%归总公司；作者仍为 50%。",
        "用户自建自用策略：平台不收分成，收益归用户本人；未确认、失败、撤销均不产生分配。",
    ])

    heading(doc, "10", "核心亮点与竞争优势")
    add_numbered(doc, [
        "权限不是简单的角色菜单，而是角色 + 组织归属链 + 接口二次校验三层控制。",
        "运营后台与运维后台共用数据但职责分离，避免业务人员接触密钥、系统开关和安全策略。",
        "用户策略从创建、回测、提交到总部双人审核形成闭环，适合治理策略市场。",
        "AI 模型支持平台统一配置和用户自带兼容端点，便于成本分层与供应商替换。",
        "行情、新闻、交易账户、策略、会员、通知和组织运营在同一客户视图中汇总，减少跨系统割裂。",
        "收入分成通过事件与分配账本记录，可按组织、代理和策略作者追溯。",
        "平台紧急停止、客户自助停单、审计日志和强制通知形成基础安全闭环。",
    ])

    heading(doc, "11", "本地验收结果")
    add_status_table(doc, [
        ("npm test", "通过", "完整构建通过，2/2 页面渲染测试通过"),
        ("目标变更 ESLint", "通过", "本次新增/修改的运维、市场源、月度分红与审批文件单独检查为 0 错误"),
        ("全仓库 ESLint", "未全绿", "旧模块仍有 84 个错误与 10 个告警，已列入改进路线图"),
        ("后台入口", "通过", "总公司超级管理员进入后先显示运营/运维两个大卡片"),
        ("运维菜单", "通过", "策略审核第一，9 个板块可访问；页面切换自动回顶部"),
        ("响应式", "通过", "1189px 浏览器宽度下无横向溢出"),
        ("行情", "部分通过", "Binance 报价/K 线成功；Coinbase 在当前网络超时"),
        ("新闻", "通过", "RSS 最新新闻成功显示"),
    ])
    add_callout(doc, "发布状态", "本次只修改和测试本地 WOX；没有推送 GitHub，没有上传 Cloudflare，也没有部署到 www.tzxsea.com。", GOLD)

    heading(doc, "A", "系统规模附录")
    add_table(doc, ["指标", "数量 / 说明"], [
        ("API 路由", "82 个 route.ts"),
        ("数据库表", "39 个 sqliteTable 定义"),
        ("本地产品图标", "1,907 个构建时生成图标"),
        ("核心角色", "9 类：hq_admin、hq_support、branch_admin、manager、supervisor、employee、customer、finance、auditor"),
        ("主要语言", "简体中文、繁体中文、英语、日语、韩语、西班牙语、俄语"),
        ("本次新增持久配置", "系统、功能、计费、集成、安全 5 个设置域"),
    ], [5.0, 12.5])
    doc.add_paragraph("文档结束。有关风险、差距、优先级和后续执行计划，请参阅单独的《AgentNovas 改进事项与下一阶段路线图》。")

    path = OUTPUT / "AgentNovas网站功能与后台体系说明书.docx"
    doc.save(path)
    return path


def build_roadmap() -> Path:
    doc = setup_document("AgentNovas 改进事项与下一阶段路线图", "逻辑不足、风险优先级、执行计划与验收标准")
    add_cover(doc, "改进事项与下一阶段路线图", "逻辑不足 · 风险优先级 · 0–90 天执行计划 · 上线验收清单", "建议：先完成资金与权限安全闭环，再扩大市场与自动交易能力")
    add_toc(doc, [
        ("01", "管理层结论"), ("02", "P0 上线阻断项"), ("03", "P1 重要增强项"),
        ("04", "P2 工程与体验优化"), ("05", "0–90 天路线图"), ("06", "分阶段验收标准"),
        ("07", "上线前检查清单"), ("08", "需要业务方最终确认的规则"),
    ])

    heading(doc, "01", "管理层结论")
    add_callout(doc, "总体判断", "网站已经具备完整产品外形和较清晰的组织、策略、运营、运维骨架，但“能演示”与“可承载真实资金/自动交易”之间仍有明显距离。下一阶段应围绕支付、交易、权限、通知、数据授权、监控与发布六条生产链路逐项闭环。", GOLD)
    add_table(doc, ["级别", "定义", "建议处理时间", "主要主题"], [
        ("P0", "不完成就不应开放真实资金或自动交易", "0–30 天", "支付对账、实盘路由、独立技术权限、服务端开关、通知、备份发布"),
        ("P1", "影响稳定运营、扩张与合规", "31–60 天", "亚洲股票数据、多语言、可观测性、API SLA、权限导航、安全强化"),
        ("P2", "影响长期研发效率和体验", "61–90 天", "代码拆分、测试体系、设计系统、无障碍、lint 技术债"),
    ], [2.0, 7.0, 3.0, 5.5])

    heading(doc, "02", "P0 上线阻断项", "下列事项建议在对外宣传“真实自动交易”或接收真实资金之前全部完成。")
    add_table(doc, ["编号", "问题 / 当前证据", "改进动作", "完成标准"], [
        ("P0-01", "会员支付仍存在演示式地址、倒计时和人工状态", "接入唯一订单地址、链上确认、回调验签、到账入账、退款和对账任务", "同一订单不重复入账；异常可追溯；每日自动对账"),
        ("P0-02", "真实交易路由默认未开放，平台 AI 周期主要为验证/演示", "建立交易执行服务、幂等订单键、状态机、撤单重试、仓位核验和交易所沙箱", "沙箱连续 7 天无重复单；断线恢复和紧急停止通过演练"),
        ("P0-03", "运维权限暂挂在 hq_admin", "拆分 hq_technical_admin、release_manager、security_admin 和只读 auditor", "技术人员无权修改业务资金；业务人员无权读取密钥"),
        ("P0-04", "部分功能开关只隐藏前端入口", "为每个高风险/收费/API 模块增加统一服务端 feature guard", "关闭后直接访问 API 也返回明确的 403/maintenance 状态"),
        ("P0-05", "Telegram/WhatsApp/邮件目前有队列与偏好，但外部发送闭环未确认", "接入供应商、重试、死信、状态回执、模板版本和告警", "强制通知在测试环境 99.9%送达并可查回执"),
        ("P0-06", "自动任务密钥本地缺失，发布/回滚/备份仍偏人工", "建立 CI/CD、环境密钥托管、迁移前备份、版本回滚和发布审计", "一次按钮/流水线完成可回滚发布；恢复演练有记录"),
        ("P0-07", "首页和部分看板把演示 KPI 与真实数据混合", "所有数字标明实时/模拟/暂无数据；生产环境禁用假延迟、假在线数和硬编码价格", "验收页面不出现无法追溯来源的业务指标"),
        ("P0-08", "平台涉及金融交易与客户敏感数据", "完成适用地区的法律、牌照、KYC/AML、隐私、数据跨境和风险披露评审", "形成律师/合规签署的上线范围与禁用地区清单"),
    ], [1.2, 5.2, 6.2, 4.9], 7.5)

    subsection(doc, "P0 推荐实施顺序")
    add_numbered(doc, [
        "先冻结真实资金与真实自动下单开关，保留模拟和只读行情。",
        "完成专用技术角色和密钥隔离，再接入 CI/CD 与备份回滚。",
        "支付与交易分别建设状态机、幂等、审计和对账，不把逻辑写在页面组件中。",
        "接通强制通知和系统告警，让资金/交易异常可以被及时发现。",
        "最后进行沙箱、故障注入、权限越权、重复回调和断线恢复验收。",
    ])

    heading(doc, "03", "P1 重要增强项")
    add_table(doc, ["编号", "现状与不足", "建议", "验收指标"], [
        ("P1-01", "韩国、日本、香港、中国股票实时行情未接入", "采购合规行情供应商，统一代码映射、交易日历、复权、币种、时区与 K 线聚合", "目标市场报价/K 线可用率 ≥99.9%，许可范围有文档"),
        ("P1-02", "公开行情源会受地域和限频影响；Coinbase 本地已超时", "服务端缓存、自动熔断、供应商优先级、健康分、速率配额和降级提示", "单源故障 30 秒内切换；客户端不出现长时间空白"),
        ("P1-03", "多语言依赖手工字典与 DOM 替换，动态内容可能漏翻", "迁移到 message catalog，按页面/组件分包，引入术语库、复数和日期货币本地化", "7 种语言无混杂；核心金融术语通过人工复核"),
        ("P1-04", "运营菜单与服务端权限并非处处完全一致", "由同一权限定义生成前端导航和 API guard；加入角色快照测试", "五级账号逐项验收，无“看得到但点不开”"),
        ("P1-05", "限流目前主要是总开关，IP 名单为精确字符串", "接入边缘限流、CIDR、可信代理头、设备指纹、登录风控和逐步阻断", "暴力登录与 API 滥用测试可自动阻断并留痕"),
        ("P1-06", "运维状态以基础计数和健康接口为主", "增加指标、结构化日志、追踪、任务队列、错误聚合、SLO 和告警", "核心接口 p95、错误率、队列延迟和供应商状态可视化"),
        ("P1-07", "AI 已支持自定义端点，但缺少 token 账单与提示词版本", "增加模型路由策略、token 用量账本、预算、提示词版本/回滚和输出评估", "可按用户/模型/功能查看日周月成本并设置预算"),
        ("P1-08", "客户登录、设备、联系方式属于敏感数据", "字段级脱敏、按需解密、查看审计、导出水印、保留期限与删除流程", "敏感字段每次查看均有审计记录，越权测试为 0"),
    ], [1.2, 5.1, 6.3, 4.9], 7.5)

    heading(doc, "04", "P2 工程与体验优化")
    add_table(doc, ["主题", "问题", "建议"], [
        ("页面架构", "app/page.tsx 体积大且包含大量内联组件，维护和测试风险高", "按路由/领域拆为首页、运营、行情、会员、通知等独立模块，并抽离服务层"),
        ("旧功能下线", "团队任务、月度目标、逾期应收等入口虽隐藏，旧接口/组件仍存在", "做调用扫描、数据保留方案和正式 deprecation，确认无依赖后再安全删除"),
        ("Lint 技术债", "全仓库当前 84 个错误、10 个告警；本次目标文件已单独清零", "按目录建立基线，先修未定义组件和无障碍，再逐步开启 CI 阻断"),
        ("测试覆盖", "自动测试目前以构建和 2 个页面渲染测试为主", "增加单元、API 集成、角色权限、收入分配、浏览器 E2E 与视觉回归"),
        ("设计系统", "部分页面仍是历史样式与新运维样式并存", "统一间距、字体、表单、状态、空态、错误和响应式断点"),
        ("可访问性", "旧表单存在 label 关联、图片和键盘操作告警", "以 WCAG 2.2 AA 为目标修复语义、焦点、对比度和屏幕阅读器"),
        ("文档与变更", "业务规则散落在聊天、代码与页面文案中", "把角色矩阵、分成公式、接口版本、数据字典纳入版本化产品文档"),
    ], [3.0, 7.0, 7.5])

    heading(doc, "05", "0–90 天路线图")
    add_table(doc, ["阶段", "时间", "主要交付", "建议负责人"], [
        ("阶段 A：安全基线", "第 1–2 周", "冻结生产交易、拆分技术角色、服务端开关、密钥托管、备份与恢复演练", "技术负责人 + 安全"),
        ("阶段 B：资金闭环", "第 3–4 周", "支付订单、链上确认、入账、退款、对账、强制通知", "后端 + 财务 + 合规"),
        ("阶段 C：交易闭环", "第 5–6 周", "沙箱执行、订单状态机、幂等、撤单、持仓核验、紧急停止演练", "交易系统 + 风控"),
        ("阶段 D：数据与运营", "第 7–8 周", "行情熔断、亚洲股票供应商、SLO、告警、角色导航一致性", "数据工程 + 运维"),
        ("阶段 E：国际化与质量", "第 9–10 周", "i18n 重构、术语审校、权限/API/收入规则测试、敏感数据审计", "前端 + QA + 法务"),
        ("阶段 F：工程收口", "第 11–12 周", "拆分单体页面、清理旧功能、降低 lint、视觉回归、上线演练与签署", "全栈 + QA + 产品"),
    ], [3.4, 2.4, 7.8, 3.9], 8.0)
    subsection(doc, "每周固定治理节奏")
    add_bullets(doc, [
        "周一：检查安全、交易、支付和供应商 SLO；确认本周发布范围。",
        "周三：执行权限回归、账本对账和异常演练；审查高风险配置变更。",
        "周五：发布候选版本、灰度验证、回滚演练；更新风险清单和业务文档。",
    ])

    heading(doc, "06", "分阶段验收标准")
    add_table(doc, ["领域", "最低验收标准"], [
        ("权限", "五级运营账号 + 4 类职能角色逐页逐接口测试；任何越权请求必须 403 并记录审计"),
        ("支付", "订单、地址、确认数、回调、重复回调、超时、退款、对账全部有状态机和幂等测试"),
        ("交易", "沙箱下单/撤单/部分成交/断线/重连/重复请求/紧急停止全部通过，持仓与交易所一致"),
        ("行情", "来源、时间戳、交易对、时区、复权规则明确；多源健康与降级告警可验证"),
        ("AI", "密钥不回显；模型超时与失败可降级；成本可归集；模型输出无法绕过硬风控"),
        ("通知", "强制通知不可关闭；失败重试和死信可查；敏感内容不明文泄露"),
        ("数据", "客户详情、登录记录、组织链和收入分配按最小权限展示；导出和查看有日志"),
        ("发布", "迁移前自动备份；发布可灰度；回滚在规定时间内完成；版本和责任人可审计"),
    ], [3.5, 14.0])

    heading(doc, "07", "上线前检查清单")
    checklist = [
        "[ ] 真实资金开关、自动交易开关默认关闭，并经过双人审批才能开启",
        "[ ] 生产、测试、开发数据库和密钥完全隔离",
        "[ ] 管理员强制 MFA；技术、财务、审核角色完成最小权限验证",
        "[ ] 支付、退款、收入分配、结算和策略作者分成全部完成对账测试",
        "[ ] 交易订单幂等、仓位核验、断线恢复和全局紧急停止演练通过",
        "[ ] 行情、新闻、股票数据源已确认商业许可、配额和 SLA",
        "[ ] 客户隐私、KYC/AML、风险披露、服务条款与禁用地区完成法律审核",
        "[ ] 监控、日志、追踪、告警、值班、事故响应和备份恢复流程已演练",
        "[ ] 7 种语言核心页面人工审校，无混杂语言和金融术语歧义",
        "[ ] 全站无演示价格、假在线数、假延迟或未标注的模拟数据",
        "[ ] 关键 API、权限、收入规则、支付与交易 E2E 自动测试进入 CI",
        "[ ] 发布版本、数据库迁移、回滚包、负责人和上线窗口已确认",
    ]
    add_bullets(doc, checklist)

    heading(doc, "08", "需要业务方最终确认的规则", "这些项目不是技术实现能自行决定的事项，应在正式上线前形成书面业务决议并版本化。")
    add_table(doc, ["待确认事项", "当前代码口径", "需要确认"], [
        ("会员收入", "毛收入总部 60% / 分公司 40%；无有效归属总部 100%", "是否适用于所有套餐、优惠、退款与税费场景"),
        ("总部部门分配", "网站可分配收益：技术 2.5%、招商 2.5%、运营 15%", "是否需要增加财务、客服、风控或预留池"),
        ("盈利分成", "每周正收益收费；20%/19%/18%/16%；亏损不结转", "是否需要高水位、亏损追回、结算时区和最低收费门槛"),
        ("个人代理", "月度业绩阶梯 20%–50%，每月清零", "业绩定义、退款冲减、跨月订单和离职结算"),
        ("策略广场", "平台 50% / 作者 50%；平台份额再按总部/分公司分配", "税费、退款、争议、盗版、作者封禁和下架后的处理"),
        ("双人审批", "策略需 2 名总部审核；运营申请由分公司最终审批且申请人不能自批", "是否要求所有资金类申请也必须两名不同审批人"),
        ("数据保留", "审计和历史记录倾向保留", "客户关闭账号、隐私删除请求与监管保留期限如何平衡"),
    ], [4.1, 7.0, 6.4], 8.0)
    add_callout(doc, "建议的下一次决策会议", "先只讨论 P0：生产支付范围、真实交易开放边界、技术/财务角色、双人审批范围、适用国家和数据许可。上述五项确认后，再确定 90 天排期和预算。", BLUE)
    doc.add_paragraph("文档结束。本路线图与《AgentNovas 网站功能与后台体系说明书》配套使用。")

    path = OUTPUT / "AgentNovas改进事项与下一阶段路线图.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    first = build_website_guide()
    second = build_roadmap()
    print(first)
    print(second)
