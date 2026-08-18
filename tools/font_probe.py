from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

names = [
    "Heiti SC", "Heiti SC Medium", "STHeiti", "STHeitiSC-Medium",
    "Arial Unicode MS", "ArialUnicodeMS", "Hiragino Sans GB W3",
    "HiraginoSansGB-W3", "AppleGothic", "Songti SC", "Microsoft YaHei",
    "HYQiHei-55J", "HYQiHei-55J Regular", "ShuS-SC",
]

doc = Document()
for name in names:
    p = doc.add_paragraph()
    r = p.add_run(f"{name} — 中文字体测试：网站功能与后台体系说明书")
    r.font.name = name
    r.font.size = Pt(18)
    fonts = r._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
doc.save("artifacts/font-probe.docx")
